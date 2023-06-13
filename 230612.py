'''

from bs4 import BeautifulSoup

html_doc = open('C:/Users/kwonohsem/Downloads/cookbook_c-2/sample-html.html')
soup = BeautifulSoup(html_doc, 'html.parser')

print('\n\nHTML이 제거된 전체 텍스트 :')
print(soup.get_text())

print("<title> 태그에 액세스 :", end = " ")
print(soup.title)

print(" <H1> 태그의 텍스트에 액세스 :", end = ' ')
print(soup.h1.string)
print("<H1> 태그의 속성에 액세스 :", end = ' ')
print(soup.img['alt'])

print('\n 존재하는 모든 <p> 태그에 액세스 :')
for p in soup.find_all('p'):
    print(p.string)

'''


'''
import feedparser

myFeed = feedparser.parse("http://feeds.mashable.com/Mashable")
print("피드 제목 : ", myFeed['feed']['title'])
print("포스트 수 : ", len(myFeed.entries))

post = myFeed.entries[0]
print("포스트 제목 :", post.title)

content = myFeed.entries[0].content[0].value #해결이 안됨..
print("콘텐츠 원본 :\n", content) #해결이 안돼.....
'''
'''
import os
from nltk.corpus.reader.plaintext import PlaintextCorpusReader



import docx
def getTextWord(wordFileName):
    doc = docx.Document(wordFileName)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)

    return '\n'.join(fullText)
docName = "C:/Users/kwonohsem/Downloads/cookbook_c-2/sample-one-line.docx"
print("Document in full : \n", getTextWord(docName))

doc = docx.Document(docName)
print("단락 수 :", len(doc.paragraphs))
print("2번 단락 :", doc.paragraphs[1].text)
print("2번 단락 스타일 :", doc.paragraphs[1].style)

print("Paragraph 1:", doc.paragraphs[0].text)
print("Number of runs in paragraph 1:", len(doc.paragraphs[0].runs))
for idx, run in enumerate(doc.paragraphs[0].runs):
    print("Run %s : %s" %(idx, run.text))

print("is Run 0 underlined: ", doc.paragraphs[0].runs[5].underline)
print("is Run 2 bold: ", doc.paragraphs[0].runs[1].bold)
print("is Run 7 italic:", doc.paragraphs[0].runs[3].italic)

from PyPDF2 import PdfReader
def getTextPDF(pdfFileName, password = ''):
    pdf_file = open(pdfFileName, 'rb')
    read_pdf = PdfReader(pdf_file)

    if password != '':
        read_pdf.decrypt(password)
    text = []
    for i in range(0, len(read_pdf.pages) ):
        text.append(read_pdf.pages[i].extract_text())
        return '\n'.join(text)

def getText(txtFileName):
    file = open(txtFileName, 'r')
    return file.read()

newCorpusDir = 'C:/Users/kwonohsem/Downloads/mycorpus/'
if not os.path.isdir(newCorpusDir):
    os.mkdir(newCorpusDir)

txt1 = getText('C:/Users/kwonohsem/Downloads/sample_feed.txt')
txt2 = getTextPDF("C:/Users/kwonohsem/Downloads/cookbook_c-2/sample-pdf.pdf")
txt3 = getTextWord("C:/Users/kwonohsem/Downloads/cookbook_c-2/sample-one-line.docx")

files = [txt1,txt2,txt3]
for idx, f in enumerate(files):
    with open(newCorpusDir + str(idx) + ' .txt', 'w') as fout:
        fout.write(f)

newCorpus = PlaintextCorpusReader(newCorpusDir, '.*')

print(newCorpus.words())
print(newCorpus.sents(newCorpus.fileids()[1]))
print(newCorpus.paras(newCorpus.fileids()[0]))


pdfFile = 'C:/Users/kwonohsem/Downloads/sample-one-line.pdf'
pdfFileEncrypted = 'C:/Users/kwonohsem/Downloads/cookbook_c-2/sample-one-line.protected.pdf'
print("PDF 1: \n", getTextPDF(pdfFile))
print("PDF 2: \n", getTextPDF(pdfFileEncrypted, 'tuffy'))

'''
'''
str = "NLTK Dolly Python"
print('다음의 인덱스에서 끝나는 부분 문자열: NLTK')
print("당므의 인덱스에서 시작하는 부분 문자열:", str[11:])
print("부분 문자열", str[5:10])
print("복잡한 방식의 부분 문자열", str[-12:-7])
if "NLTK" in str:
    print("NLTK를 찾았습니다")

replaced = str.replace("Dolly", "Dorothy")
print("대체된 문자열:", replaced)
print("각 문자(character) 액세스:")
for s in replaced:
    print(s)
'''
'''

namesList = ['유나', '지은', '스튜어트', '케빈']
sentence = '우리 강아지는 소파 위에서 잔다'

names = ';'.join(namesList)
print(type(names), ':', names)

wordList = sentence.split(' ')
print((type(wordList)), ':', wordList)
additionExample = "파이썬" + "파이썬" + "파이썬"
multiplicationExample = "파이썬" * 2
print("텍스트 덧셈 :", additionExample)
print("텍스트 곱셈 :", multiplicationExample)

str = "Python NLTK"
print(str[1])
print(str[-3])
'''

#----------------------
'''

from nltk.corpus import wordnet as wn
type = 'n'
synsets = wn.all_synsets(type)

lemmas = []
for synset in synsets:
    for lemma in synset.lemmas():
        lemmas.append(lemma.name())

lemmas = set(lemmas)
count = 0
for lemma in lemmas:
    count = count + len(wn.synsets(lemma, type))

print('개별 기본형 합계: ', len(lemmas))
print("총 뜻: ", count)
print(type, '(명사)의 다의어 평균: ', count/len(lemmas))
'''
'''

from nltk.corpus import wordnet as wn
woman = wn.synset('woman.n.01')
bed = wn.synset('bed.n.01')
print(woman.hypernyms())
woman_paths = woman.hypernym_paths()

for idx, path in enumerate(woman_paths):
    print('\n\n 상위어 경로 :', idx + 1)
    for synset in path:
        print(synset.name(), ', ', end = ' ')

types_of_beds = bed.hyponyms()
print('\n\nbed의 형태(하위어):', types_of_beds)
print(sorted(set(lemma.name() for synset in types_of_beds for lemma in synset.lemmas())))
'''
'''

from nltk.corpus import wordnet as wn
chair = 'chair'
chair_synsets = wn.synsets(chair)
print('의자(Chair)의 뜻 Synsets : ', chair_synsets, '\n\n')

for synset in chair_synsets:
    print(synset, ': ')
    print("Definition: ", synset.definition())
    print("Lemmas/Synonymous words: ", synset.lemma_names())
    print("Example: ", synset.examples(), '\n')
'''
'''

import nltk, matplotlib
from nltk.corpus import webtext
print(webtext.fileids())

wbt_words = webtext.words('singles.txt')
fdist = nltk.FreqDist(wbt_words)

print('최대 발생 토큰 "', fdist.max(), '" 수 : ', fdist[fdist.max()])
print('말뭉치 내 총 고유 토큰 수 : ', fdist.N())
print("말뭉치에서 가장 흔한 10개 단어는 다음과 같습니다.")
print(fdist.most_common(10))

print('개인 광고의 빈도 분포')
print(fdist.tabulate())

fdist.plot(cumulative = True)
'''

'''

import nltk
from nltk.corpus import brown

print(brown.categories())

genres = ['fiction', 'humor', 'romance']
whwords = ['what', 'which', 'how', 'why', 'when', 'where', 'who']

for i in range(0, len(genres)):
    genre = genres[i]
    print()
    print("'" + genre + "' wh단어 분석")
    genre_text = brown.words(categories = genre)
    fdist = nltk.FreqDist(genre_text)

    for wh in whwords:
        print(wh + ':', fdist[wh], end = ' ')


'''


'''
from nltk.corpus import CategorizedPlaintextCorpusReader
reader = CategorizedPlaintextCorpusReader(r"C:/Users/kwonohsem/Desktop/mix20_rand700_tokens_cleaned/tokens", r".*/.txt", cat_pattern = r"(\w+)/*")
print(reader.categories())
print(reader.fileids())

posFiles = reader.fileids(categories = 'pos')
negFiles = reader.fileids(categories = 'neg')

print(posFiles)
print(negFiles)

#posFiles = reader.fileids(categories = 'tokens')
#print(posFiles)

from random import randint
fileP = posFiles[randint(0, len(posFiles) - 1)]
fileN = negFiles[randint(0, len(negFiles) - 1)]
print(fileP)
print(fileN)

for w in reader.words(fileP):
    print(w + ' ', end = '')
    if (w == '.'):
        print()

for w in reader.words(fileN):
    print(w + ' ', end = '')
    if (w == '.'):
        print()
'''
'''

from nltk.corpus import reuters

files = reuters.fileids()
print(files)

words16097 = reuters.words(['test/16097'])
print(words16097)

x = 20

words20 = reuters.words(['test/16097'])[:20]
print(words20)

words20 = reuters.words(['test/16097'])[:x]
print(words20)


reutersGenres = reuters.categories()
print(reutersGenres)

y = ['acq', 'alum', 'barley', 'bop', 'carcass', 'castor-oil', 'cocoa', 'coconut', 'coconut-oil', 'coffee', 'copper', 'copra-cake', 'corn', 'cotton', 'cotton-oil', 'cpi', 'cpu', 'crude', 'dfl', 'dlr', 'dmk', 'earn', 'fuel', 'gas', 'gnp', 'gold', 'grain', 'groundnut', 'groundnut-oil', 'heat', 'hog', 'housing', 'income', 'instal-debt', 'interest', 'ipi', 'iron-steel', 'jet', 'jobs', 'l-cattle', 'lead', 'lei', 'lin-oil', 'livestock', 'lumber', 'meal-feed', 'money-fx', 'money-supply', 'naphtha', 'nat-gas', 'nickel', 'nkr', 'nzdlr', 'oat', 'oilseed', 'orange', 'palladium', 'palm-oil', 'palmkernel', 'pet-chem', 'platinum', 'potato', 'propane', 'rand', 'rape-oil', 'rapeseed', 'reserves', 'retail', 'rice', 'rubber', 'rye', 'ship', 'silver', 'sorghum', 'soy-meal', 'soy-oil', 'soybean', 'strategic-metal', 'sugar', 'sun-meal', 'sun-oil', 'sunseed', 'tea', 'tin', 'trade', 'veg-oil', 'wheat', 'wpi', 'yen', 'zinc']
print(len(y))

for w in reuters.words(categories=['bop', 'cocoa']):
    print(w+' ', end = '')
    if(w == '.'):
        print()


'''
